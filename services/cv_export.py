"""
CV Export Service — Phase 3E

Generates a polished .docx that weaves the base CV with tailored headline,
match narrative, and talking-points for a specific target job.
"""
from __future__ import annotations

import io
import logging
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# Brand colour — matches the UI accent (#10D9A0)
ACCENT = RGBColor(0x10, 0xD9, 0xA0)
MUTED  = RGBColor(0x64, 0x74, 0x8B)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _section_heading(doc: Document, text: str) -> None:
    """Uppercase small-caps section divider with accent bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)

    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = ACCENT

    # Bottom border
    pPr    = p._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "10D9A0")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _body(doc: Document, text: str, size: float = 10) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.size = Pt(size)


def _bullet(doc: Document, text: str, size: float = 10) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_tailored_cv_docx(
    cv_data:  dict[str, Any],
    tailored: dict[str, Any],
    job:      dict[str, Any],
) -> bytes:
    """
    Build and return a .docx as bytes.
    Combines the candidate's parsed CV with tailored sections for the target job.
    """
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Name ─────────────────────────────────────────────────────────────────
    name = cv_data.get("name") or "Candidate"
    h    = doc.add_heading(name, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h.runs:
        h.runs[0].font.color.rgb = ACCENT

    # Tailored headline (italic, centred)
    if tailored.get("tailored_headline"):
        hl = doc.add_paragraph()
        hl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hl.add_run(tailored["tailored_headline"])
        run.font.size   = Pt(11)
        run.font.italic = True

    # Contact line
    contact_parts = [
        v for v in [cv_data.get("email"), cv_data.get("phone"), cv_data.get("location")]
        if v
    ]
    if contact_parts:
        cp = doc.add_paragraph(" · ".join(contact_parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cp.runs:
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].font.color.rgb = MUTED

    # ── Tailored section ──────────────────────────────────────────────────────
    job_label = f"{job.get('title', '')} @ {job.get('company', '')}".strip(" @")
    _section_heading(doc, f"Tailored for: {job_label}")

    if tailored.get("match_narrative"):
        _body(doc, tailored["match_narrative"])

    if tailored.get("cover_points"):
        for pt in tailored["cover_points"]:
            _bullet(doc, pt)

    # Skills to lead with
    if tailored.get("top_skills_to_highlight"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run("Lead with: ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(" · ".join(s.title() for s in tailored["top_skills_to_highlight"]))
        run2.font.size = Pt(10)
        run2.font.color.rgb = ACCENT

    # ── Professional Summary ──────────────────────────────────────────────────
    if cv_data.get("summary"):
        _section_heading(doc, "Professional Summary")
        _body(doc, cv_data["summary"])

    # ── Skills ────────────────────────────────────────────────────────────────
    all_skills = cv_data.get("skills", [])
    if all_skills:
        _section_heading(doc, "Skills")
        _body(doc, " · ".join(s.title() for s in all_skills))

    # ── Experience ────────────────────────────────────────────────────────────
    experience = cv_data.get("experience", [])
    if experience:
        _section_heading(doc, "Experience")
        for exp in experience:
            # Role title + company
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
            run.bold = True
            run.font.size = Pt(10)

            # Date range
            start = exp.get("start_date", "")
            end   = exp.get("end_date",   "")
            date_str = f"{start} – {end}".strip(" –")
            if date_str:
                dr = p.add_run(f"   {date_str}")
                dr.font.size = Pt(9)
                dr.font.color.rgb = MUTED

            # Description
            if exp.get("description"):
                _body(doc, exp["description"])

    # ── Education ─────────────────────────────────────────────────────────────
    education = cv_data.get("education", [])
    if education:
        _section_heading(doc, "Education")
        for edu in education:
            degree = edu.get("degree", "")
            field  = edu.get("field", "")
            inst   = edu.get("institution", "")
            year   = edu.get("year", "")
            p = doc.add_paragraph()
            run1 = p.add_run(f"{degree} in {field}".strip(" in"))
            run1.bold = True
            run1.font.size = Pt(10)
            run2 = p.add_run(f"  —  {inst}  ({year})".replace("  ()", ""))
            run2.font.size = Pt(10)

    # ── Skill gaps note (small print) ─────────────────────────────────────────
    gaps = tailored.get("skill_gaps", [])
    if gaps:
        _section_heading(doc, "Areas to Develop")
        _body(doc, f"Highlight in interview: {', '.join(gaps)}", size=9)

    # ── Serialise ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
