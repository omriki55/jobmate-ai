"""
Cover Letter Service — generates personalized cover letters for job applications.

Uses Claude to create a tailored 3-paragraph cover letter matching
the candidate's CV to a specific job posting.
"""
from __future__ import annotations

import io
import json
import logging
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from config.settings import ANTHROPIC_API_KEY
from services.cv_export import _section_heading, _body, ACCENT, MUTED

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Prompt
# ---------------------------------------------------------------------------

COVER_LETTER_PROMPT = """\
You are an expert career coach writing a compelling cover letter.

## Candidate Profile
Name: {name}
Domain: {primary_domain}
Seniority: {seniority_level}
Years of experience: {total_years_experience}
Skills: {skills}
Professional summary: {summary}

## Target Role
Title: {job_title}
Company: {company}
Location: {location}
Description: {description}
Requirements: {requirements}

## Task
Write a personalized 3-paragraph cover letter and return ONLY a valid JSON object:
{{
  "cover_letter": "Full cover letter text (3 paragraphs separated by \\n\\n)",
  "key_points": ["Key selling point 1", "Key selling point 2", "Key selling point 3"],
  "tone": "professional|enthusiastic|confident"
}}

Rules:
- Paragraph 1: Hook — why this role at this company excites the candidate.
- Paragraph 2: Evidence — 2-3 specific achievements/skills that match requirements.
- Paragraph 3: Closing — enthusiasm + call to action.
- Be specific — reference the company name, role title, and real skills.
- Keep it under 300 words total.
- Return ONLY the JSON object.
"""


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

async def generate_cover_letter(
    cv_data: dict[str, Any],
    job_data: dict[str, Any],
) -> dict[str, Any]:
    """Generate a personalized cover letter for a job application."""
    if not ANTHROPIC_API_KEY or ANTHROPIC_API_KEY.startswith("your_"):
        return _fallback_cover_letter(cv_data, job_data)

    try:
        from anthropic import AsyncAnthropic
        client = AsyncAnthropic(api_key=ANTHROPIC_API_KEY)

        prompt = COVER_LETTER_PROMPT.format(
            name=cv_data.get("name", "Candidate"),
            primary_domain=cv_data.get("primary_domain", "your field"),
            seniority_level=cv_data.get("seniority_level", "mid"),
            total_years_experience=cv_data.get("total_years_experience", 0),
            skills=", ".join(cv_data.get("skills", [])[:15]),
            summary=(cv_data.get("summary", "") or "")[:400],
            job_title=job_data.get("title", ""),
            company=job_data.get("company", ""),
            location=job_data.get("location", ""),
            description=(job_data.get("description", "") or "")[:800],
            requirements=", ".join(job_data.get("requirements", [])[:15]),
        )

        msg = await client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
            raw = raw.rsplit("```", 1)[0]
        return json.loads(raw.strip())

    except Exception as exc:
        logger.warning("Cover letter generation failed: %s", exc)
        return _fallback_cover_letter(cv_data, job_data)


# ---------------------------------------------------------------------------
# Fallback
# ---------------------------------------------------------------------------

def _fallback_cover_letter(
    cv_data: dict[str, Any],
    job_data: dict[str, Any],
) -> dict[str, Any]:
    """Template-based cover letter when Claude is unavailable."""
    name = cv_data.get("name", "Candidate")
    domain = cv_data.get("primary_domain", "my field")
    seniority = cv_data.get("seniority_level", "experienced")
    yoe = cv_data.get("total_years_experience", 0)
    skills = cv_data.get("skills", [])[:5]
    title = job_data.get("title", "the open position")
    company = job_data.get("company", "your company")

    skills_str = ", ".join(skills[:3]) if skills else domain

    para1 = (
        f"I am writing to express my strong interest in the {title} position "
        f"at {company}. With {yoe} years of experience in {domain}, I am "
        f"confident that my background and skills make me an excellent fit "
        f"for this role."
    )
    para2 = (
        f"Throughout my career, I have developed deep expertise in {skills_str}. "
        f"As a {seniority}-level professional, I have consistently delivered "
        f"results that drive business impact. My experience aligns well with "
        f"the requirements outlined in your job description, and I am eager "
        f"to bring this expertise to {company}."
    )
    para3 = (
        f"I am excited about the opportunity to contribute to {company}'s "
        f"continued success and would welcome the chance to discuss how my "
        f"background can add value to your team. Thank you for considering "
        f"my application."
    )

    return {
        "cover_letter": f"{para1}\n\n{para2}\n\n{para3}",
        "key_points": [
            f"{yoe} years of {domain} experience",
            f"Expertise in {skills_str}",
            f"{seniority.title()}-level professional with proven track record",
        ],
        "tone": "professional",
    }


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

def generate_cover_letter_docx(
    cv_data: dict[str, Any],
    job: dict[str, Any],
    cl_data: dict[str, Any],
) -> bytes:
    """Build a .docx cover letter and return as bytes."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Candidate name
    name = cv_data.get("name") or "Candidate"
    h = doc.add_heading(name, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h.runs:
        h.runs[0].font.color.rgb = ACCENT

    # Contact line
    contact_parts = [
        v for v in [cv_data.get("email"), cv_data.get("phone"), cv_data.get("location")]
        if v
    ]
    if contact_parts:
        cp = doc.add_paragraph(" · ".join(contact_parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cp.runs:
            cp.runs[0].font.size = Pt(9)
            cp.runs[0].font.color.rgb = MUTED

    # Target role heading
    job_label = f"{job.get('title', '')} @ {job.get('company', '')}".strip(" @")
    _section_heading(doc, f"Cover Letter: {job_label}")

    # Cover letter paragraphs
    letter_text = cl_data.get("cover_letter", "")
    for para in letter_text.split("\n\n"):
        para = para.strip()
        if para:
            _body(doc, para)

    # Sign-off
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    run = p.add_run(f"Sincerely,\n{name}")
    run.font.size = Pt(10)

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
